import pandas as pd
from docx import Document
import os
from tqdm import tqdm  # <- Barra de progresso

# Carregar os dados
acompanhamento = pd.read_csv("acompanhamento.csv", encoding="latin1", dtype=str)
racks = pd.read_csv("quantidade-de-racks.csv", encoding="latin1", dtype=str)
enderecos = pd.read_csv("dados-redes.csv", encoding="latin1", dtype=str)

# Limpar nomes das colunas
acompanhamento.columns = acompanhamento.columns.str.strip()
racks.columns = racks.columns.str.strip()
enderecos.columns = enderecos.columns.str.strip()

# Criar pasta de saída principal
saida_principal = "saida_docs"
os.makedirs(saida_principal, exist_ok=True)

# Iterar com barra de progresso
for _, row in tqdm(acompanhamento.iterrows(), total=len(acompanhamento), desc="Gerando documentos"):
    codigo_escola = row["CIE"].strip()
    quantidade_swits = int(row["Tota SW"].strip())

    # Procurar quantidade de racks com base no CIE
    match_rack = racks[racks["CIE"].str.strip() == codigo_escola]
    quantidade_racks = int(match_rack.iloc[0]["RACKS"].strip()) if not match_rack.empty else 0

    # Procurar dados da escola
    match_endereco = enderecos[enderecos["CIE"].str.strip() == codigo_escola]
    if not match_endereco.empty:
        dados = match_endereco.iloc[0]
    else:
        dados = pd.Series(dtype=str)

    # Criar pasta para a escola
    pasta_escola = os.path.join(saida_principal, codigo_escola)
    os.makedirs(pasta_escola, exist_ok=True)

    doc = Document()

    
    # Conteúdo principal
    doc.add_heading(f"Termo de Entregáveis - CIE: {codigo_escola}", level=1)
    doc.add_paragraph("Instruções para o técnico: As evidências são obrigatórias. É necessário enviar uma foto para cada equipamento atendido, conforme os requisitos abaixo:")


    # Cabeçalho com dados da escola
    doc.add_heading(f"Escola: {dados.get('NOME_COMPLETO_ESCOLA', 'NÃO INFORMADO')}", level=1)
    doc.add_paragraph(f"Diretoria de Ensino: {dados.get('DIRETORIA DE ENSINO', 'NÃO INFORMADO')}")
    doc.add_paragraph(f"Município: {dados.get('MUNICÍPIO', 'NÃO INFORMADO')}")
    doc.add_paragraph(f"Endereço: {dados.get('ENDEREÇO', 'NÃO INFORMADO')}")
    doc.add_paragraph(f"CEP: {dados.get('CEP', 'NÃO INFORMADO')}")
    doc.add_paragraph(f"Quantidade de Switches para instalar: {dados.get('QUANTIDADE DE SW PARA INSTALAR', quantidade_swits)}")
    doc.add_paragraph("")

    # Endereçamento
    doc.add_heading("Endereçamento IP dos Switches:", level=2)
    for i in range(1, 10):
        campo = f"ENDEREÇAMENTO {i}"
        ip = dados.get(campo)
        if pd.notna(ip):
            doc.add_paragraph(f"{campo}: {ip}")

    doc.add_paragraph(f"MÁSCARA: {dados.get('MÁSCARA', 'NÃO INFORMADO')}")
    doc.add_paragraph(f"GATEWAY: {dados.get('GATEWAY', 'NÃO INFORMADO')}")

    
    doc.add_paragraph(f"{quantidade_racks} fotos dos racks antes da substituição.")
    doc.add_paragraph(f"{quantidade_swits} fotos dos switches depois de terem sido substituídos.")

    doc.add_heading("Rat (Relatório de Atendimento Técnico):", level=2)
    doc.add_paragraph(f"{quantidade_swits} RAT com os números de série dos switches devolvidos.")

    doc.add_heading("Fotos dos switches devolvidos:", level=2)
    doc.add_paragraph(f"{quantidade_swits} fotos dos switches que foram devolvidos.")

    doc.add_heading("(Display version) Versão do Sistema Operacional/Firmware:", level=2)
    doc.add_paragraph(f"{quantidade_swits} capturas de tela mostrando a versão do sistema.")

    doc.add_heading("(Display device manuinfo) Print da tela:", level=2)
    doc.add_paragraph(f"{quantidade_swits} Print da tela dos dispositivos contendo número de série, data de fabricação e nome do fornecedor.")

    doc.add_heading("Teste de Ping:(Print)", level=2)
    doc.add_paragraph(f"{quantidade_swits} capturas de tela do teste de ping para 8.8.8.8.")

    doc.add_heading("(Display cloud State) Conexão de Rede Estabelecida:", level=2)
    doc.add_paragraph(f"{quantidade_swits} capturas de tela demonstrando conexão de rede ativa (termo 'estabelecido').")

    # Salvar dentro da pasta da escola
    nome_arquivo = os.path.join(pasta_escola, f"Termo_{codigo_escola}.docx")
    doc.save(nome_arquivo)
